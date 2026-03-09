"""
Document Processor - Handles file reading, parsing, and chunking
"""
import os
import re
from pathlib import Path
from typing import Optional, Generator
from dataclasses import dataclass, field
from enum import Enum
import xxhash

from pypdf import PdfReader
from docx import Document as DocxDocument
from striprtf.striprtf import rtf_to_text
from bs4 import BeautifulSoup
from langchain_text_splitters import RecursiveCharacterTextSplitter
from dotenv import load_dotenv
from rich.console import Console

load_dotenv()
console = Console()


class FileType(Enum):
    """Supported file types"""
    CODE = "code"
    MARKDOWN = "markdown"
    DOCUMENT = "document"
    DATA = "data"
    WEB = "web"
    UNKNOWN = "unknown"


# File type mappings
FILE_TYPE_MAP = {
    # Code
    ".ts": FileType.CODE,
    ".tsx": FileType.CODE,
    ".py": FileType.CODE,
    ".js": FileType.CODE,
    ".css": FileType.CODE,
    # Markdown
    ".md": FileType.MARKDOWN,
    # Documents
    ".txt": FileType.DOCUMENT,
    ".rtf": FileType.DOCUMENT,
    ".pdf": FileType.DOCUMENT,
    ".docx": FileType.DOCUMENT,
    # Data
    ".csv": FileType.DATA,
    ".json": FileType.DATA,
    ".jsonc": FileType.DATA,
    ".yaml": FileType.DATA,
    ".yml": FileType.DATA,
    # Web
    ".html": FileType.WEB,
    ".htm": FileType.WEB,
}


@dataclass
class DocumentChunk:
    """A chunk of processed document"""
    content: str
    chunk_index: int
    total_chunks: int
    metadata: dict = field(default_factory=dict)


@dataclass
class ProcessedDocument:
    """Result of document processing"""
    file_path: str
    file_hash: str
    file_type: FileType
    raw_content: str
    chunks: list[DocumentChunk]
    metadata: dict = field(default_factory=dict)
    error: Optional[str] = None


class DocumentProcessor:
    """Handles reading, parsing, and chunking documents"""
    
    def __init__(
        self,
        chunk_size: int = None,
        chunk_overlap: int = None,
        include_extensions: list[str] = None
    ):
        self.chunk_size = chunk_size or int(os.getenv("CHUNK_SIZE", "512"))
        self.chunk_overlap = chunk_overlap or int(os.getenv("CHUNK_OVERLAP", "50"))
        
        ext_str = os.getenv("INCLUDE_EXTENSIONS", ".ts,.tsx,.py,.md,.txt")
        self.include_extensions = include_extensions or ext_str.split(",")
        
        # Initialize text splitters for different content types
        self._code_splitter = RecursiveCharacterTextSplitter(
            chunk_size=self.chunk_size,
            chunk_overlap=self.chunk_overlap,
            separators=["\n\nclass ", "\n\ndef ", "\n\nasync def ", "\nfunction ", "\n\n", "\n", " "],
        )
        
        self._prose_splitter = RecursiveCharacterTextSplitter(
            chunk_size=self.chunk_size,
            chunk_overlap=self.chunk_overlap,
            separators=["\n\n## ", "\n\n### ", "\n\n", "\n", ". ", " "],
        )
        
        self._data_splitter = RecursiveCharacterTextSplitter(
            chunk_size=self.chunk_size,
            chunk_overlap=self.chunk_overlap,
            separators=["\n\n", "\n", ",", " "],
        )
    
    def is_supported(self, file_path: str | Path) -> bool:
        """Check if file type is supported"""
        path = Path(file_path)
        return path.suffix.lower() in self.include_extensions
    
    def get_file_type(self, file_path: str | Path) -> FileType:
        """Determine file type from extension"""
        path = Path(file_path)
        return FILE_TYPE_MAP.get(path.suffix.lower(), FileType.UNKNOWN)
    
    def compute_hash(self, content: str) -> str:
        """Compute fast hash of content"""
        return xxhash.xxh64(content.encode()).hexdigest()
    
    def read_file(self, file_path: str | Path) -> tuple[Optional[str], Optional[str]]:
        """Read file content based on type. Returns (content, error)"""
        path = Path(file_path)
        
        if not path.exists():
            return None, f"File not found: {path}"
        
        suffix = path.suffix.lower()
        
        try:
            # PDF
            if suffix == ".pdf":
                reader = PdfReader(path)
                content = "\n\n".join(page.extract_text() or "" for page in reader.pages)
                return content, None
            
            # DOCX
            if suffix == ".docx":
                doc = DocxDocument(path)
                content = "\n\n".join(para.text for para in doc.paragraphs)
                return content, None
            
            # RTF
            if suffix == ".rtf":
                with open(path, "r", encoding="utf-8", errors="ignore") as f:
                    rtf_content = f.read()
                content = rtf_to_text(rtf_content)
                return content, None
            
            # HTML
            if suffix in (".html", ".htm"):
                with open(path, "r", encoding="utf-8", errors="ignore") as f:
                    html_content = f.read()
                soup = BeautifulSoup(html_content, "lxml")
                # Remove script and style elements
                for element in soup(["script", "style"]):
                    element.decompose()
                content = soup.get_text(separator="\n", strip=True)
                return content, None
            
            # Plain text / code (default)
            with open(path, "r", encoding="utf-8", errors="ignore") as f:
                content = f.read()
            return content, None
            
        except Exception as e:
            return None, f"Error reading {path}: {str(e)}"
    
    def chunk_content(self, content: str, file_type: FileType) -> list[str]:
        """Split content into chunks based on file type"""
        if not content or not content.strip():
            return []
        
        if file_type == FileType.CODE:
            return self._code_splitter.split_text(content)
        elif file_type in (FileType.DATA,):
            return self._data_splitter.split_text(content)
        else:
            return self._prose_splitter.split_text(content)
    
    def process(self, file_path: str | Path) -> ProcessedDocument:
        """Process a single file into chunks"""
        path = Path(file_path)
        file_type = self.get_file_type(path)
        
        # Read content
        content, error = self.read_file(path)
        
        if error or not content:
            return ProcessedDocument(
                file_path=str(path.absolute()),
                file_hash="",
                file_type=file_type,
                raw_content="",
                chunks=[],
                error=error or "Empty content"
            )
        
        # Compute hash
        file_hash = self.compute_hash(content)
        
        # Chunk content
        chunk_texts = self.chunk_content(content, file_type)
        
        # Build chunk objects with metadata
        chunks = [
            DocumentChunk(
                content=chunk_text,
                chunk_index=i,
                total_chunks=len(chunk_texts),
                metadata={
                    "file_path": str(path.absolute()),
                    "file_name": path.name,
                    "file_type": file_type.value,
                    "chunk_index": i,
                    "total_chunks": len(chunk_texts),
                }
            )
            for i, chunk_text in enumerate(chunk_texts)
        ]
        
        return ProcessedDocument(
            file_path=str(path.absolute()),
            file_hash=file_hash,
            file_type=file_type,
            raw_content=content,
            chunks=chunks,
            metadata={
                "file_name": path.name,
                "file_extension": path.suffix,
                "file_size": path.stat().st_size,
                "file_type": file_type.value,
                "chunk_count": len(chunks),
            }
        )
    
    def process_directory(
        self,
        directory: str | Path,
        recursive: bool = True
    ) -> Generator[ProcessedDocument, None, None]:
        """Process all supported files in a directory"""
        path = Path(directory)
        
        if not path.exists():
            console.print(f"[red]✗[/red] Directory not found: {path}")
            return
        
        pattern = "**/*" if recursive else "*"
        
        for file_path in path.glob(pattern):
            if file_path.is_file() and self.is_supported(file_path):
                yield self.process(file_path)


# Quick test
def test_processor():
    processor = DocumentProcessor()
    
    # Test with a sample file
    test_content = """
# Test Document

This is a test markdown document for chunking.

## Section 1

Some content in section 1 that we want to process and chunk appropriately.

## Section 2

More content here with different information.
    """
    
    # Simulate processing
    chunks = processor.chunk_content(test_content, FileType.MARKDOWN)
    console.print(f"[blue]Chunks created:[/blue] {len(chunks)}")
    for i, chunk in enumerate(chunks):
        console.print(f"[dim]Chunk {i}:[/dim] {chunk[:50]}...")


if __name__ == "__main__":
    test_processor()
